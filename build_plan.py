"""
PageGains Homepage Redesign — Implementation Plan Generator
Uses python-docx to produce a structured Word document.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page setup ───────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = section.right_margin = Inches(1)
section.top_margin  = section.bottom_margin = Inches(1)

# ── Colours ──────────────────────────────────────────────────────────────────
RED   = RGBColor(0xFF, 0x4D, 0x3A)
INK   = RGBColor(0x0B, 0x0B, 0x0C)
GREY  = RGBColor(0x6B, 0x6B, 0x70)
CODE_BG = "F0EDEA"

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_run(run, bold=False, italic=False, size=None, color=None, font="Arial"):
    run.bold   = bold
    run.italic = italic
    if size:  run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    run.font.name = font

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.keep_with_next = True
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top  = OxmlElement('w:top')
    top.set(qn('w:val'), 'single'); top.set(qn('w:sz'), '12')
    top.set(qn('w:space'), '6');    top.set(qn('w:color'), 'FF4D3A')
    pBdr.append(top); pPr.append(pBdr)
    r = p.add_run(text); set_run(r, bold=True, size=18, color=RED)
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text); set_run(r, bold=True, size=13, color=INK)
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text); set_run(r, bold=True, size=11, color=RED)
    return p

def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); set_run(r, size=10, color=INK)
    return p

def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.2)
    r = p.add_run(text); set_run(r, italic=True, size=9.5, color=GREY)
    return p

def label(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text.upper()); set_run(r, bold=True, size=8.5, color=RED, font="Courier New")
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.2)
    r = p.add_run(text); set_run(r, size=10, color=INK)
    return p

def kv(key, val):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(key + ": "); set_run(r1, bold=True, size=10, color=INK)
    r2 = p.add_run(val);        set_run(r2, size=10, color=GREY)

def code_block(lines):
    def shade(p):
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), CODE_BG); pPr.append(shd)
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Inches(0.15)
        shade(p)
        r = p.add_run(line if line else " ")
        r.font.name = "Courier New"; r.font.size = Pt(8.5); r.font.color.rgb = INK
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(6)

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '4')
    bot.set(qn('w:space'), '1');    bot.set(qn('w:color'), 'EBE8E1')
    pBdr.append(bot); pPr.append(pBdr)

# ─────────────────────────────────────────────────────────────────────────────
# COVER
# ─────────────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(36); p.paragraph_format.space_after = Pt(6)
r = p.add_run("PAGEGAINS HOMEPAGE REDESIGN"); set_run(r, bold=True, size=24, color=RED)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER; p2.paragraph_format.space_after = Pt(4)
r2 = p2.add_run("Full Implementation Plan"); set_run(r2, size=14, color=INK)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER; p3.paragraph_format.space_after = Pt(8)
r3 = p3.add_run("Vite + React  |  GSAP + ScrollTrigger  |  Lenis.js  |  Frontend Prototype")
set_run(r3, italic=True, size=10, color=GREY, font="Courier New")

divider()

# ─────────────────────────────────────────────────────────────────────────────
# 1. PROJECT CONTEXT
# ─────────────────────────────────────────────────────────────────────────────
h1("1. PROJECT CONTEXT")
body("PageGains is a CRO audit tool. Paste a URL, get 10-20 ranked conversion fixes in ~60 "
     "seconds. First audit is $3.99 with a money-back guarantee, no subscription.")
body("This is a frontend-only prototype spec. No backend integration. Goal: a show-stopping "
     "redesign to pitch to the site owner as a commission. Do not fabricate any statistics, "
     "testimonials, pricing, or copy. All approved data is in Section 3.")
kv("Stack", "Vite + React (no router -- single page)")
kv("Animation", "GSAP + ScrollTrigger")
kv("Smooth scroll", "Lenis.js")
kv("Fonts", "Geist, Instrument Serif, JetBrains Mono (Google Fonts)")
kv("Target", "Desktop-first, fully responsive to mobile")
divider()

# ─────────────────────────────────────────────────────────────────────────────
# 2. DESIGN DIRECTION
# ─────────────────────────────────────────────────────────────────────────────
h1("2. DESIGN DIRECTION")

h2("Style: Editorial Brutalism")
body("Push the existing editorial newspaper aesthetic into asymmetric, grid-breaking brutalist "
     "territory. The current site is symmetrical, polite, and safe. The redesign should feel "
     "heavy, show-stopping, and unapologetic.")

label("Core Adjectives")
for adj in ["Asymmetrical", "Blunt", "Heavy", "Show-stopping", "Urgent"]:
    bullet(adj)

label("Core Rules")
bullet("Typography IS the layout in key sections -- the headline is the structure, not content inside a box.")
bullet("Hard 2px solid ink-colour rules as section dividers, not just whitespace.")
bullet("1-2 deliberate grid breaks per section: bleeds, rotations, oversized elements.")
bullet("Grain texture: SVG noise overlay at ~4% opacity on paper background for warmth and depth.")
bullet("Depth via hard offset shadows -- push beyond brand defaults (10-12px on key elements).")
bullet("Horizontal scroll in Testimonials section ONLY.")
bullet("Marquee ticker strip as momentum interrupt between Hero and How It Works.")
bullet("NEVER use transition-all. Animate only transform and opacity.")
bullet("No glassmorphism. No purple gradients. No generic illustration packs.")

h2("Colour")
body("Keep the existing palette exactly. Coral red #FF4D3A is brand primary. "
     "Paper #F6F4EF is page background. Ink #0B0B0C is body text and dark UI. "
     "Crimson #C41E2A is for urgency signals only (live indicators, critical labels).")

h2("Typography")
kv("Display", "Geist 700-800 -- headlines, large pricing figures")
kv("Accent", "Instrument Serif italic -- 3-6 words per section max, for texture")
kv("UI / Meta", "JetBrains Mono -- eyebrows, stats, form chrome, section labels")
body("Sizes may be pushed beyond brand token defaults for more visual impact.")

h2("Signature Elements")
bullet("Skewed red highlight chip on 1-2 words per major headline. CSS: transform: skew(-4deg), box-shadow: 6px 6px 0 var(--pg-ink).")
bullet("Hard offset shadows on CTAs: 6px 6px 0 var(--pg-ink). On URL bar: 10px 10px 0 var(--pg-red).")
bullet("Mono uppercase eyebrows with letter-spacing 0.14em.")
divider()

# ─────────────────────────────────────────────────────────────────────────────
# 3. PAGE RHYTHM + REAL DATA
# ─────────────────────────────────────────────────────────────────────────────
h1("3. PAGE RHYTHM + REAL DATA REFERENCE")

h2("Light / Dark Section Map")
note("Apply data-theme=\"dark\" attribute directly to the <section> element. tokens.css handles the rest.")
rows = [
    ("Nav",           "LIGHT"),
    ("Hero",          "LIGHT"),
    ("Marquee Strip", "RED (#FF4D3A background, 2px ink borders top/bottom)"),
    ("How It Works",  "LIGHT"),
    ("Sample Output", "LIGHT"),
    ("Testimonials",  "DARK  <-- data-theme=\"dark\""),
    ("Pricing",       "DARK  <-- continues (separate element, same theme)"),
    ("FAQ",           "LIGHT <-- returns to light"),
    ("Final CTA",     "DARK  <-- data-theme=\"dark\""),
    ("Footer",        "DARK  <-- continues"),
]
for sname, theme in rows:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{sname:<20}"); set_run(r1, bold=True, size=10, color=INK, font="Courier New")
    r2 = p.add_run(theme);          set_run(r2, size=10, color=GREY, font="Courier New")

h2("Real Data -- Use Only These. No Fabrication.")
label("Pricing")
kv("Single audit", "$3.99 per audit")
kv("Free to start", "Score + issue preview are free. Full report costs $3.99 to unlock.")
kv("Starter pack", "5 audits  |  $9.99  |  $2.00/audit  |  50% off")
kv("Growth pack",  "20 audits |  $29.99 |  $1.50/audit  |  62% off")
kv("Scale pack",   "50 audits |  $49.99 |  $1.00/audit  |  75% off")

label("Speed & Volume Stats")
kv("Result time",    "~60 seconds / median 52s")
kv("Issues",         "Average 17 per page / range 10-20 ranked")
kv("Signals",        "120 signals scanned")
kv("Benchmark",      "8,000+ converting SaaS pages")
kv("Agency contrast","Agencies charge $3,000-$10,000/month and take weeks (from public FAQ)")

label("Testimonials (verbatim)")
testimonials = [
    ("Marco D., Indie hacker",
     "\"For four bucks I got a sharper audit than the agency I paid $2,400 for last quarter. Uncomfortable truths.\""),
    ("Jonathan B., Myloops.net",
     "\"Conversion rate went from 2.6% to 3.3% within a few weeks. Ranked by impact is the killer feature.\" "
     "(Photo: jonathan-b.jpg in public/)"),
    ("Eric F., Founder at Editaura",
     "\"Every single tip was extremely useful.\""),
    ("Section H2 pull quote (Eric's)",
     "\"It found things a $3k consultant missed.\""),
    ("Jonathan stat (from public FAQ)",
     "+$1,933/month for Myloops.net -- public info, can appear near Jonathan's card"),
]
for name, quote in testimonials:
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Inches(0.2)
    r1 = p.add_run(name + ": ");  set_run(r1, bold=True, size=9.5, color=INK)
    r2 = p.add_run(quote);        set_run(r2, size=9.5, color=GREY)

label("Assets -- copy to /public/")
for a in ["pagegains-logo-new.png", "audit-result-screenshot.webp",
          "jonathan-b.jpg", "gainbot-icon.png", "favicon.png"]:
    bullet(a)

label("Launch Badges (load via external URLs)")
badges = [
    ("Nick Launches",  "https://nicklaunches.com/badges/featured-dark.svg",
                       "https://nicklaunches.com/products/pagegains/"),
    ("Startup Fame",   "https://startupfa.me/badges/featured-badge.webp",
                       "https://startupfa.me/s/pagegains"),
    ("Launchit",       "https://www.launchit.site/badges/launchit-dark.svg",
                       "https://www.launchit.site/launches/pagegains"),
    ("Better Launch",  "https://www.betterlaunch.co/badge-light.svg",
                       "https://www.betterlaunch.co"),
    ("ListBulb",       "https://www.listbulb.com/featured-on-listbulb-light.svg",
                       "https://www.listbulb.com/tools/pagegains"),
]
for name, img, link in badges:
    bullet(f"{name}  |  img: {img}  |  href: {link}")
divider()

# ─────────────────────────────────────────────────────────────────────────────
# PHASE 0
# ─────────────────────────────────────────────────────────────────────────────
h1("PHASE 0 -- PROJECT SETUP")

h2("Install Commands")
code_block([
    "npm create vite@latest pagegains-redesign -- --template react",
    "cd pagegains-redesign",
    "npm install",
    "npm install gsap lenis",
    "",
    "# Copy from brand_assets/ into project:",
    "# tokens.css              -->  src/tokens.css",
    "# pagegains-logo-new.png  -->  public/",
    "# audit-result-screenshot.webp  -->  public/",
    "# jonathan-b.jpg          -->  public/",
    "# gainbot-icon.png        -->  public/",
    "# favicon.png             -->  public/",
])

h2("index.html <head> Additions")
code_block([
    '<link rel="preconnect" href="https://fonts.googleapis.com">',
    '<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>',
    '<link href="https://fonts.googleapis.com/css2?',
    '      family=Geist:wght@400;600;700;800',
    '      &family=Instrument+Serif:ital@0,1',
    '      &family=JetBrains+Mono:wght@400;500',
    '      &display=swap" rel="stylesheet">',
    '<link rel="preload" as="image" href="/audit-result-screenshot.webp">',
    '<link rel="icon" href="/favicon.png">',
    "<title>PageGains -- Find What's Killing Your Conversions In 60s</title>",
])

h2("Folder Structure")
code_block([
    "src/",
    "  components/",
    "    Eyebrow.jsx        -- JetBrains Mono label",
    "    HighlightChip.jsx  -- Skewed red keyword chip",
    "    CTAButton.jsx      -- Hard offset shadow CTA",
    "    URLBar.jsx         -- URL input / hero conversion element",
    "    Marquee.jsx        -- Horizontal scrolling ticker strip",
    "  sections/",
    "    Nav.jsx",
    "    Hero.jsx",
    "    HowItWorks.jsx",
    "    SampleOutput.jsx",
    "    Testimonials.jsx",
    "    Pricing.jsx",
    "    FAQ.jsx",
    "    FinalCTA.jsx",
    "    Footer.jsx",
    "  App.jsx",
    "  main.jsx",
    "  tokens.css           -- copied from brand_assets/tokens.css",
    "  global.css",
])

h2("main.jsx -- Lenis + GSAP Initialisation")
note("Wires Lenis into GSAP's ticker so ScrollTrigger stays in sync with smooth scroll.")
code_block([
    "import React from 'react'",
    "import ReactDOM from 'react-dom/client'",
    "import Lenis from 'lenis'",
    "import gsap from 'gsap'",
    "import ScrollTrigger from 'gsap/ScrollTrigger'",
    "import App from './App'",
    "import './tokens.css'",
    "import './global.css'",
    "",
    "gsap.registerPlugin(ScrollTrigger)",
    "",
    "const lenis = new Lenis({",
    "  duration: 1.2,",
    "  easing: (t) => Math.min(1, 1.001 - Math.pow(2, -10 * t)),",
    "  smoothWheel: true,",
    "  wheelMultiplier: 0.8,",
    "  touchMultiplier: 1.5,",
    "})",
    "",
    "lenis.on('scroll', ScrollTrigger.update)",
    "gsap.ticker.add((time) => { lenis.raf(time * 1000) })",
    "gsap.ticker.lagSmoothing(0)",
    "",
    "ReactDOM.createRoot(document.getElementById('root')).render(",
    "  <React.StrictMode><App /></React.StrictMode>",
    ")",
])

h2("App.jsx")
code_block([
    "import Nav         from './sections/Nav'",
    "import Hero        from './sections/Hero'",
    "import HowItWorks  from './sections/HowItWorks'",
    "import SampleOutput from './sections/SampleOutput'",
    "import Testimonials from './sections/Testimonials'",
    "import Pricing     from './sections/Pricing'",
    "import FAQ         from './sections/FAQ'",
    "import FinalCTA    from './sections/FinalCTA'",
    "import Footer      from './sections/Footer'",
    "",
    "export default function App() {",
    "  return (",
    "    <>",
    "      <Nav />",
    "      <main>",
    "        <Hero />",
    "        <HowItWorks />",
    "        <SampleOutput />",
    "        <Testimonials />",
    "        <Pricing />",
    "        <FAQ />",
    "        <FinalCTA />",
    "      </main>",
    "      <Footer />",
    "    </>",
    "  )",
    "}",
])
divider()

# ─────────────────────────────────────────────────────────────────────────────
# PHASE 1
# ─────────────────────────────────────────────────────────────────────────────
h1("PHASE 1 -- GLOBAL DESIGN PASS")
body("Establish the global visual language before building any section. Every later stage "
     "builds into this system.")

h2("global.css -- Base Reset & Body")
code_block([
    "*, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }",
    "html { scroll-behavior: auto; }  /* Lenis handles smooth scroll */",
    "",
    "body {",
    "  background: var(--pg-paper);",
    "  color: var(--pg-ink);",
    "  font-family: var(--pg-font-display);",
    "  font-size: var(--pg-text-body);",
    "  line-height: var(--pg-leading-body);",
    "  -webkit-font-smoothing: antialiased;",
    "  overflow-x: hidden;",
    "}",
    "",
    "/* Grain texture overlay */",
    "body::after {",
    "  content: '';",
    "  position: fixed; inset: 0; z-index: 9999;",
    "  pointer-events: none; opacity: 0.045;",
    "  background-image: url(\"data:image/svg+xml,%3Csvg viewBox='0 0 256 256'",
    "    xmlns='http://www.w3.org/2000/svg'%3E%3Cfilter id='n'%3E",
    "    %3CfeTurbulence type='fractalNoise' baseFrequency='0.9'",
    "    numOctaves='4' stitchTiles='stitch'/%3E%3C/filter%3E",
    "    %3Crect width='100%25' height='100%25' filter='url(%23n)'/%3E",
    "    %3C/svg%3E\");",
    "  background-size: 200px 200px;",
    "}",
    "",
    "::selection { background: var(--pg-red); color: #fff; }",
    "::-webkit-scrollbar { width: 6px; }",
    "::-webkit-scrollbar-track { background: var(--pg-line); }",
    "::-webkit-scrollbar-thumb { background: var(--pg-red); border-radius: 3px; }",
    "img { max-width: 100%; display: block; }",
    "a { color: inherit; text-decoration: none; }",
    "",
    ".pg-section {",
    "  padding: var(--pg-section-y) var(--pg-gutter);",
    "  max-width: var(--pg-max-width);",
    "  margin: 0 auto;",
    "}",
    "@media (max-width: 768px) {",
    "  .pg-section { padding: var(--pg-section-y-sm) var(--pg-gutter-sm); }",
    "}",
])

h2("Shared Components")

h3("Eyebrow.jsx")
body("JetBrains Mono, 11px, uppercase, letter-spacing 0.14em, color var(--pg-mute). Props: text.")
code_block([
    "export default function Eyebrow({ text }) {",
    "  return <span className='eyebrow'>{text}</span>",
    "}",
    "",
    "/* eyebrow CSS */",
    ".eyebrow {",
    "  font-family: var(--pg-font-mono);",
    "  font-size: var(--pg-text-meta);",
    "  color: var(--pg-mute);",
    "  letter-spacing: var(--pg-tracking-meta);",
    "  text-transform: uppercase;",
    "  display: block;",
    "  margin-bottom: 16px;",
    "}",
])

h3("HighlightChip.jsx")
body("Skewed red chip. Use on 1-2 words per major headline maximum. Props: children.")
code_block([
    "export default function HighlightChip({ children }) {",
    "  return <mark className='highlight-chip'>{children}</mark>",
    "}",
    "",
    ".highlight-chip {",
    "  background: var(--pg-red); color: #fff;",
    "  padding: 0 0.1em; display: inline-block;",
    "  transform: skew(-4deg);",
    "  box-shadow: 6px 6px 0 var(--pg-ink);",
    "  font-style: normal;",
    "}",
])

h3("CTAButton.jsx")
body("Hard offset shadow button. Props: label, href, size ('default'|'small'), theme ('red'|'ink').")
code_block([
    "export default function CTAButton({ label, href, size='default', theme='red' }) {",
    "  return (",
    "    <a href={href} className={`cta-btn cta-btn--${size} cta-btn--${theme}`}>",
    "      {label}",
    "    </a>",
    "  )",
    "}",
    "",
    ".cta-btn {",
    "  display: inline-flex; align-items: center;",
    "  font-family: var(--pg-font-display); font-weight: 700; font-size: 15px;",
    "  border-radius: var(--pg-radius-sm); padding: 14px 28px; cursor: pointer;",
    "  transition: transform 150ms var(--pg-ease-default),",
    "              box-shadow 150ms var(--pg-ease-default);",
    "}",
    ".cta-btn--red { background: var(--pg-red); color: #fff; box-shadow: 6px 6px 0 var(--pg-ink); }",
    ".cta-btn--ink { background: var(--pg-ink); color: #fff; box-shadow: 6px 6px 0 var(--pg-red); }",
    ".cta-btn--small { padding: 10px 20px; font-size: 13px; }",
    ".cta-btn:hover { transform: translate(3px, 3px); box-shadow: 3px 3px 0 var(--pg-ink); }",
    ".cta-btn--ink:hover { box-shadow: 3px 3px 0 var(--pg-red); }",
])

h3("URLBar.jsx")
body("Primary conversion element. Styled as a terminal prompt. No real form action needed "
     "for this prototype.")
code_block([
    "export default function URLBar() {",
    "  return (",
    "    <div className='url-bar'>",
    "      <div className='url-bar__prefix'>https://</div>",
    "      <input className='url-bar__input' type='text'",
    "        placeholder='yoursite.com/pricing'",
    "        aria-label='Enter your page URL' />",
    "      <button className='url-bar__btn'>Audit my page <arrow/></button>",
    "    </div>",
    "  )",
    "}",
    "",
    ".url-bar {",
    "  display: flex; align-items: center;",
    "  border: 2px solid var(--pg-ink);",
    "  border-radius: var(--pg-radius-sm);",
    "  background: #fff;",
    "  box-shadow: 10px 10px 0 var(--pg-red);",
    "  overflow: hidden;",
    "  transition: box-shadow 150ms var(--pg-ease-default);",
    "}",
    ".url-bar:focus-within { box-shadow: 10px 10px 0 var(--pg-ink); }",
    ".url-bar__prefix {",
    "  font-family: var(--pg-font-mono); font-size: 14px; color: var(--pg-mute);",
    "  padding: 0 12px 0 16px; border-right: 1px solid var(--pg-line);",
    "  display: flex; align-items: center; white-space: nowrap;",
    "}",
    ".url-bar__input {",
    "  flex: 1; border: none; outline: none;",
    "  font-family: var(--pg-font-mono); font-size: 14px;",
    "  padding: 16px 12px; background: transparent; color: var(--pg-ink);",
    "}",
    ".url-bar__btn {",
    "  background: var(--pg-ink); color: #fff; border: none;",
    "  font-family: var(--pg-font-display); font-weight: 700; font-size: 14px;",
    "  padding: 16px 24px; cursor: pointer; white-space: nowrap;",
    "  transition: background 150ms;",
    "}",
    ".url-bar__btn:hover { background: var(--pg-red); }",
])
divider()

# ─────────────────────────────────────────────────────────────────────────────
# PHASE 2
# ─────────────────────────────────────────────────────────────────────────────
h1("PHASE 2 -- TONE PASS")
body("Before building any section fully, stub ALL sections with placeholder headings "
     "and background colours only. Scroll the full page and verify:")
for check in [
    "Light > Red marquee > Light > Dark > Light > Dark rhythm reads correctly",
    "Lenis scrolls smoothly at duration 1.2 (tune to 0.8-1.4 if needed -- do it here)",
    "Grain texture is present but not distracting",
    "Typography scale feels massive enough in the hero stub",
    "Section heights have natural breathing room",
    "Hard offset shadows on URLBar and CTAButton feel right in context",
    "Dark sections (#0B0B0C) contrast sharply against light (#F6F4EF)",
]:
    bullet(check)
body("Only after the tone pass is confirmed, proceed to Stage 1.")
divider()

# ─────────────────────────────────────────────────────────────────────────────
# STAGE 1 -- NAV + HERO
# ─────────────────────────────────────────────────────────────────────────────
h1("STAGE 1 -- NAV + HERO")
note("Highest conversion impact. Nail this before anything else.")

h2("Nav")
kv("Position",  "fixed, top: 0, z-index: 100")
kv("Height",    "64px")
kv("Layout",    "flex, space-between, align-items center, max-width 1400px, padding 0 40px")
kv("Left",      "Logo: <img src='/pagegains-logo-new.png' height='28'> in <a href='/'>")
kv("Centre",    "Links: 'How it works' / 'Pricing' / 'FAQ' | JetBrains Mono 12px uppercase")
kv("Right",     "<CTAButton label='Audit my page ->' size='small' href='#hero-url-bar'>")
kv("On scroll", "Add class .nav--scrolled when scrollY > 20 --> adds paper background + border-bottom")
kv("Mobile",    "Hide centre links. Logo + CTA only.")
code_block([
    "useEffect(() => {",
    "  const fn = () => navRef.current.classList.toggle('nav--scrolled', window.scrollY > 20)",
    "  window.addEventListener('scroll', fn, { passive: true })",
    "  return () => window.removeEventListener('scroll', fn)",
    "}, [])",
])

h2("Hero")
body("Asymmetric two-column grid desktop (65% left / 35% right). Single column mobile.")

label("Grid CSS")
code_block([
    ".hero-grid {",
    "  display: grid;",
    "  grid-template-columns: 65fr 35fr;",
    "  gap: 60px; align-items: center;",
    "  max-width: 1400px; margin: 0 auto;",
    "  padding: 140px 40px 100px;",
    "}",
    "@media (max-width: 900px) {",
    "  .hero-grid { grid-template-columns: 1fr; padding: 100px 20px 60px; }",
    "}",
])

h3("Left Column Contents (top to bottom)")

label("1. Eyebrow with live indicator")
body("JetBrains Mono, 11px, uppercase. Text: 'CONVERSION AUDIT . v1.2 . ' + pulsing red dot "
     "(8px circle, background var(--pg-crimson), animate opacity 1<->0.3 at 1.5s) "
     "+ ' LIVE . BUILT FOR SAAS'")

label("2. H1 -- Rotating Headline")
body("Font: Geist 800, clamp(52px, 8vw, 110px), line-height 0.88, letter-spacing -0.045em.")
body("Structure: 'Your [rotating type] is [HighlightChip]losing[/HighlightChip] conversions right now.'")
body("Rotating values: 'pricing page' / 'landing page' / 'homepage' / 'signup flow' -- CSS only, no JS library.")
code_block([
    "/* CSS rotation -- one word visible at a time */",
    ".rotating-word { position: relative; display: inline-block; }",
    ".rotating-word span {",
    "  position: absolute; left: 0; top: 0;",
    "  opacity: 0; animation: wordCycle 12s linear infinite;",
    "}",
    ".rotating-word span:nth-child(1) { animation-delay: 0s;  }",
    ".rotating-word span:nth-child(2) { animation-delay: 3s;  }",
    ".rotating-word span:nth-child(3) { animation-delay: 6s;  }",
    ".rotating-word span:nth-child(4) { animation-delay: 9s;  }",
    "@keyframes wordCycle {",
    "  0%, 5%   { opacity: 0; transform: translateY(8px);  }",
    "  10%, 20% { opacity: 1; transform: translateY(0);    }",
    "  25%, 100%{ opacity: 0; transform: translateY(-8px); }",
    "}",
])

label("3. Subline")
kv("Text",   "10-20 ranked fixes. ~60 seconds. $3.99.")
kv("Font",   "Geist 400, 22px, line-height 1.35, color var(--pg-body)")
kv("Margin", "24px top, 32px bottom")

label("4. URLBar component (id='hero-url-bar')")

label("5. Stats row")
body("flex, gap: 32px, margin-top: 20px. JetBrains Mono 11px.")
body("'Median 52s'  .  'Avg 17 issues'  .  'Money-back guarantee'")
body("Numbers in var(--pg-ink) bold. Labels in var(--pg-mute).")

h3("Right Column")
body("audit-result-screenshot.webp")
code_block([
    ".hero-screenshot {",
    "  width: 100%; max-width: 520px;",
    "  transform: rotate(1.5deg);",
    "  box-shadow: var(--pg-shadow-float);",
    "  border-radius: var(--pg-radius-md);",
    "  margin-left: -20px; position: relative; z-index: 2;",
    "  will-change: transform;",
    "}",
    "@media (max-width: 900px) {",
    "  .hero-screenshot { transform: none; margin-left: 0; margin-top: 40px; }",
    "}",
])

label("GSAP -- onMount Animations (not scroll-triggered)")
code_block([
    "useEffect(() => {",
    "  const tl = gsap.timeline()",
    "  tl.from(eyebrowRef.current,    { opacity:0, y:16 }, 0.10)",
    "    .from(h1Ref.current,         { opacity:0, y:24 }, 0.20)",
    "    .from(sublineRef.current,    { opacity:0, y:16 }, 0.35)",
    "    .from(urlBarRef.current,     { opacity:0, y:12, scale:0.98 }, 0.45)",
    "    .from(statsItems,            { opacity:0, y:8, stagger:0.08 }, 0.55)",
    "    .from(screenshotRef.current, { opacity:0, x:30 }, 0.25)",
    "  // All: ease: 'power3.out', duration: 0.8",
    "}, [])",
])
divider()

# ─────────────────────────────────────────────────────────────────────────────
# STAGE 2
# ─────────────────────────────────────────────────────────────────────────────
h1("STAGE 2 -- MARQUEE + HOW IT WORKS + SAMPLE OUTPUT")

h2("Marquee Strip")
body("Full-width red band. 2px ink borders top and bottom. Between Hero and How It Works.")
code_block([
    ".marquee-strip {",
    "  background: var(--pg-red); overflow: hidden;",
    "  border-top: 2px solid var(--pg-ink);",
    "  border-bottom: 2px solid var(--pg-ink);",
    "  height: 42px; display: flex; align-items: center;",
    "}",
    ".marquee-track {",
    "  display: flex; width: max-content;",
    "  animation: marquee-scroll 35s linear infinite;",
    "  will-change: transform;",
    "}",
    ".marquee-track span {",
    "  font-family: var(--pg-font-mono); font-size: 11px; color: #fff;",
    "  text-transform: uppercase; letter-spacing: 0.12em;",
    "  white-space: nowrap; padding: 0 32px;",
    "}",
    "@keyframes marquee-scroll {",
    "  from { transform: translateX(0); }",
    "  to   { transform: translateX(-50%); }",
    "}",
])
body("Marquee text (paste twice for seamless loop):")
note("YOUR PRICING PAGE IS BLEEDING SIGNUPS . $3.99 . 52 SECONDS . "
     "YOUR HOMEPAGE IS COSTING YOU DEMOS . 10-20 RANKED FIXES . "
     "NO SUBSCRIPTION . MONEY-BACK GUARANTEE . ")

label("Scroll velocity speed-up (wire in main.jsx or Marquee useEffect)")
code_block([
    "lenis.on('scroll', ({ velocity }) => {",
    "  const speed = Math.max(0.5, 1 + Math.abs(velocity) * 0.3)",
    "  marqueeTrack.style.animationDuration = `${35 / speed}s`",
    "})",
])

h2("How It Works")
kv("Eyebrow", "HOW IT WORKS . 3 STEPS")
kv("H2",      "Three clicks. A [Instrument Serif italic]brutal[/italic], honest report.")
kv("H2 font", "Geist 800, clamp(42px, 5.4vw, 78px), letter-spacing -0.035em")

label("Grid -- 2px gap = the border")
code_block([
    ".steps-grid {",
    "  display: grid;",
    "  grid-template-columns: repeat(3, 1fr);",
    "  gap: 2px;",
    "  background: var(--pg-line);  /* gap colour IS the border */",
    "  margin-top: 60px;",
    "}",
    ".step-card { background: var(--pg-paper); padding: 40px 36px; }",
    "@media (max-width: 768px) {",
    "  .steps-grid { grid-template-columns: 1fr; }",
    "}",
])

label("Step 01 / PASTE")
kv("Eyebrow", "01 / PASTE")
kv("Number",  "'01' JetBrains Mono 64px, color var(--pg-line-strong), mb 16px -- decorative")
kv("Title",   "Drop in any URL")
kv("Body",    "Landing, product, pricing, checkout. We crawl the page the way a real visitor sees it -- above-the-fold, scroll depth, forms, CTAs.")

label("Step 02 / ANALYZE")
kv("Title",   "AI scans 120 signals")
kv("Body",    "Clarity, friction, trust, hierarchy, readability, pricing anchors, form length, mobile layout. Benchmarked against 8,000+ converting SaaS pages.")

label("Step 03 / FIX")
kv("Title",   "Ranked fixes, not fluff.")
kv("Body",    "Every issue comes with the exact rewrite, severity, and a predicted lift range. Ship it in an afternoon.")

label("GSAP ScrollTrigger -- Staggered cards")
code_block([
    "gsap.from(cardEls, {",
    "  opacity: 0, y: 40, duration: 0.7, ease: 'power3.out', stagger: 0.15,",
    "  scrollTrigger: { trigger: gridRef.current, start: 'top 80%', once: true }",
    "})",
])

h2("Sample Output")
kv("Eyebrow", "SAMPLE OUTPUT . REAL PAGE")
kv("H2",      "This is what [$3.99 in HighlightChip] actually buys you.")
body("Image: audit-result-screenshot.webp. Max-width: 1100px, centered, margin-top: 48px.")
code_block([
    ".sample-screenshot {",
    "  width: 100%; max-width: 1100px; margin: 48px auto 0;",
    "  border-radius: var(--pg-radius-md);",
    "  border: 1px solid var(--pg-line);",
    "  box-shadow: var(--pg-shadow-float);",
    "  transform: perspective(1200px) rotateX(8deg);",
    "  will-change: transform; backface-visibility: hidden;",
    "}",
])
label("GSAP Scrub -- Screenshot 'stands up' as user scrolls in")
code_block([
    "gsap.fromTo(screenshotRef.current,",
    "  { rotateX: 10, opacity: 0.7 },",
    "  {",
    "    rotateX: 0, opacity: 1, ease: 'none',",
    "    scrollTrigger: {",
    "      trigger: screenshotRef.current,",
    "      start: 'top 85%', end: 'top 30%',",
    "      scrub: 1.5,",
    "    }",
    "  }",
    ")",
])
divider()

# ─────────────────────────────────────────────────────────────────────────────
# STAGE 3
# ─────────────────────────────────────────────────────────────────────────────
h1("STAGE 3 -- TESTIMONIALS + PRICING")

h2("Testimonials")
body("DARK section. data-theme='dark' on <section>. Background var(--pg-ink).")
kv("Eyebrow",      "WHAT FOUNDERS SAY")
kv("H2 pull quote","\"It found things a $3k consultant missed.\"")
kv("Attribution",  "-- Eric F., Founder at Editaura | JetBrains Mono 12px below H2")
kv("H2 font",      "Instrument Serif italic at clamp(36px, 4vw, 64px) -- creates nice tension with display font")

label("Horizontal Scroll Container")
code_block([
    ".testimonials-track-wrapper {",
    "  overflow-x: scroll; cursor: grab;",
    "  -ms-overflow-style: none; scrollbar-width: none;",
    "  margin-top: 60px;",
    "}",
    ".testimonials-track-wrapper::-webkit-scrollbar { display: none; }",
    ".testimonials-track-wrapper:active { cursor: grabbing; }",
    ".testimonials-track { display: flex; gap: 24px; padding: 8px 40px 40px; }",
    "",
    ".testimonial-card {",
    "  background: var(--pg-neutral-900);",
    "  border: 1px solid var(--pg-neutral-800);",
    "  border-radius: var(--pg-radius-xl);",
    "  padding: 36px; min-width: 380px; flex-shrink: 0;",
    "  box-shadow: 8px 8px 0 var(--pg-red);",
    "  transition: transform 200ms var(--pg-ease-default);",
    "}",
    ".testimonial-card:hover { transform: translate(-3px, -3px); }",
])

label("Drag-to-scroll (useEffect)")
code_block([
    "useEffect(() => {",
    "  const el = trackRef.current",
    "  let isDown = false, startX, scrollLeft",
    "  const onDown  = e => { isDown = true; startX = e.pageX - el.offsetLeft; scrollLeft = el.scrollLeft }",
    "  const onUp    = () => isDown = false",
    "  const onMove  = e => {",
    "    if (!isDown) return",
    "    e.preventDefault()",
    "    el.scrollLeft = scrollLeft - (e.pageX - el.offsetLeft - startX) * 1.5",
    "  }",
    "  el.addEventListener('mousedown', onDown)",
    "  el.addEventListener('mouseleave', onUp)",
    "  el.addEventListener('mouseup', onUp)",
    "  el.addEventListener('mousemove', onMove)",
    "  return () => {",
    "    el.removeEventListener('mousedown', onDown)",
    "    el.removeEventListener('mouseleave', onUp)",
    "    el.removeEventListener('mouseup', onUp)",
    "    el.removeEventListener('mousemove', onMove)",
    "  }",
    "}, [])",
])

label("Card 1 -- Marco D. (leads -- strongest agency contrast)")
bullet("Stars: ***** in var(--pg-red)")
bullet("Quote (18px, line-height 1.4): \"For four bucks I got a sharper audit than the agency I paid $2,400 for last quarter. Uncomfortable truths.\"")
bullet("Attribution: Marco D. . Indie hacker -- JetBrains Mono 11px, var(--pg-neutral-500)")

label("Card 2 -- Jonathan B. (only hard CR numbers)")
bullet("Stars: *****")
bullet("Stat callout ABOVE quote: '+$1,933/mo' -- Geist 800, 40px, color var(--pg-red). 'Myloops.net' below in mono 11px.")
bullet("Quote: \"Conversion rate went from 2.6% to 3.3% within a few weeks. Ranked by impact is the killer feature.\"")
bullet("Attribution: jonathan-b.jpg (32px circle) + Jonathan B. . Myloops.net")

label("Card 3 -- Eric F.")
bullet("Stars: *****")
bullet("Quote: \"Every single tip was extremely useful.\"")
bullet("Attribution: Eric F. . Founder, Editaura")

h2("Pricing")
body("DARK section. Continues from Testimonials. border-top: 1px solid var(--pg-neutral-800).")
kv("Eyebrow", "PRICING . ONE PRICE, ONE THING")
kv("H2",      "No subscription. No calls. [Instrument Serif italic]No plans.[/italic]")

label("Hero Audit Card")
code_block([
    ".pricing-hero-card {",
    "  max-width: 480px; margin: 48px auto 0;",
    "  background: var(--pg-neutral-900);",
    "  border: 2px solid var(--pg-neutral-700);",
    "  border-radius: var(--pg-radius-xl);",
    "  box-shadow: 10px 10px 0 var(--pg-red);",
    "  padding: 48px;",
    "}",
])
label("Card Contents (top to bottom)")
bullet("'$3.99' -- Geist 800, 80px, color var(--pg-red), letter-spacing -0.04em. Animated via counter on scroll-enter (see Stage 5).")
bullet("Label: 'Per audit . no subscription' -- JetBrains Mono 11px, var(--pg-neutral-500)")
bullet("Divider: 1px solid var(--pg-neutral-800), margin 24px 0")
for feat in [
    "Full conversion audit of one page",
    "Ranked list of 10-20 issues with predicted lift",
    "Exact rewrites for headlines, CTAs, microcopy",
    "10x Chat with GainBot (ask anything about any fix)",
    "Sharable HTML + PDF report",
    "Results in under 60 seconds",
]:
    bullet(f"[->] {feat}", level=1)
bullet("<CTAButton label='Start for free ->' theme='red'> -- full width")
bullet("Fine print: 'Money-back guarantee . Audits never expire . No subscription' -- mono 10px, centered")

label("Pack Cards Row")
body("Three cards below hero card. flex, gap: 16px, margin-top: 48px.")
code_block([
    ".pack-card {",
    "  flex: 1; background: var(--pg-neutral-900);",
    "  border: 1px solid var(--pg-neutral-800);",
    "  border-radius: var(--pg-radius-lg); padding: 28px 24px;",
    "  transition: box-shadow 200ms var(--pg-ease-default),",
    "              transform 200ms var(--pg-ease-default);",
    "}",
    ".pack-card:hover { box-shadow: 6px 6px 0 var(--pg-red); transform: translate(-3px, -3px); }",
    ".pack-card--popular { border-color: var(--pg-red); }",
    "/* POPULAR badge: mono 9px uppercase, bg var(--pg-red), color #fff, border-radius pill */",
])
kv("Starter", "5 audits . $9.99 . $2.00/audit . 50% off")
kv("Growth",  "20 audits . $29.99 . $1.50/audit . 62% off -- POPULAR badge")
kv("Scale",   "50 audits . $49.99 . $1.00/audit . 75% off")
divider()

# ─────────────────────────────────────────────────────────────────────────────
# STAGE 4
# ─────────────────────────────────────────────────────────────────────────────
h1("STAGE 4 -- FAQ + FINAL CTA + FOOTER")

h2("FAQ")
body("Returns to LIGHT. Narrow column for reading: max-width 800px, margin 0 auto.")
kv("Eyebrow", "FAQ")
kv("H2",      "Things you'll want to know.")
kv("Open item","border-left: 3px solid var(--pg-red), padding-left: 16px")

label("Accordion height animation (GSAP, not CSS transition)")
code_block([
    "function toggleItem(index) {",
    "  const answer = answerRefs.current[index]",
    "  const isOpen = openIndex === index",
    "  gsap.to(answer, {",
    "    height: isOpen ? 0 : answer.scrollHeight,",
    "    duration: 0.3, ease: 'power2.inOut',",
    "    onComplete: () => { if (!isOpen) answer.style.height = 'auto' }",
    "  })",
    "  setOpenIndex(isOpen ? null : index)",
    "}",
])

label("Question priority order (get answers verbatim from original/index.html)")
for i, q in enumerate([
    "How much does it cost?",
    "How fast is it?",
    "Do I need to create an account?",
    "What pages should I audit first?",
    "What kind of fixes will I get?",
    "Does it work for SaaS landing pages?",
    "Is my data secure?",
    "What if the analysis results aren't helpful?",
], 1):
    bullet(f"{i}. {q}")

h2("Final CTA")
body("DARK section. Centered. Minimal. The closing argument.")
kv("Padding",   "120px 40px")
kv("H2 line 1", "'Stop guessing.' -- Instrument Serif italic, Geist 800, clamp(48px, 6vw, 90px)")
kv("H2 line 2", "'Paste the URL.' -- Geist 800, same size")
kv("Subline",   "'Free to start . $3.99 to unlock . Money-back guarantee' -- mono 12px, var(--pg-neutral-500), margin 24px 0 40px")
kv("URLBar",    "Centered, max-width 640px, margin 0 auto")

label("GSAP ScrollTrigger")
code_block([
    "gsap.from(h2El, {",
    "  opacity: 0, y: 40, scale: 0.97,",
    "  duration: 0.8, ease: 'power3.out',",
    "  scrollTrigger: { trigger: h2El, start: 'top 80%', once: true }",
    "})",
])

h2("Footer")
body("DARK section. border-top: 1px solid var(--pg-neutral-800). Padding: 60px 40px 40px.")
label("Badges row")
body("flex, gap: 16px, flex-wrap: wrap, justify-content: center. Badge height: 32px.")
for name, img, link in badges:
    bullet(f"{name}  |  <img src='{img}' height='32'>  |  href: {link}")
label("Footer bottom")
bullet("Left: logo (24px height) + '(c) 2026 PageGains . J2FB SAS' -- mono 11px, var(--pg-neutral-600)")
bullet("Right: Blog . Free Resources . Legal . Contact -- mono 11px, gap 24px, hover: var(--pg-neutral-300)")
divider()

# ─────────────────────────────────────────────────────────────────────────────
# STAGE 5
# ─────────────────────────────────────────────────────────────────────────────
h1("STAGE 5 -- MOTION LAYER")
body("Apply animations across all sections. Only transform and opacity -- no layout animations.")

h2("Default Scroll Reveal")
body("Add className='reveal' to: all Eyebrow components, all H2s, lead paragraphs, Final CTA button.")
code_block([
    "gsap.utils.toArray('.reveal').forEach(el => {",
    "  gsap.from(el, {",
    "    opacity: 0, y: 30, duration: 0.7, ease: 'power3.out',",
    "    scrollTrigger: { trigger: el, start: 'top 82%', once: true }",
    "  })",
    "})",
])

h2("Staggered Grid Reveals")
body("Apply to: step cards, pack cards, FAQ items.")
code_block([
    "gsap.from(cardEls, {",
    "  opacity: 0, y: 40, duration: 0.7, ease: 'power3.out', stagger: 0.12,",
    "  scrollTrigger: { trigger: containerEl, start: 'top 78%', once: true }",
    "})",
])

h2("Pricing Counter ($0.00 -> $3.99)")
code_block([
    "const counter = { val: 0 }",
    "ScrollTrigger.create({",
    "  trigger: priceRef.current, start: 'top 80%', once: true,",
    "  onEnter: () => {",
    "    gsap.to(counter, {",
    "      val: 3.99, duration: 1.2, ease: 'power2.out',",
    "      snap: { val: 0.01 },",
    "      onUpdate: () => {",
    "        priceRef.current.textContent = '$' + counter.val.toFixed(2)",
    "      }",
    "    })",
    "  }",
    "})",
])

h2("Hero Screenshot Parallax")
code_block([
    "gsap.to(screenshotRef.current, {",
    "  y: -50, ease: 'none',",
    "  scrollTrigger: {",
    "    trigger: heroRef.current,",
    "    start: 'top top', end: 'bottom top',",
    "    scrub: true,",
    "  }",
    "})",
])

h2("Motion Hard Rules")
for rule in [
    "NEVER use transition-all. Write each property: transition: transform 150ms, box-shadow 150ms",
    "ONLY animate transform and opacity (GPU-accelerated)",
    "will-change: transform on all GSAP-animated elements",
    "backface-visibility: hidden on all 3D transforms (rotateX in Sample Output)",
    "All one-shot scroll animations: once: true",
    "Wrap all GSAP in a prefers-reduced-motion check:",
]:
    bullet(rule)
code_block([
    "const pRM = window.matchMedia('(prefers-reduced-motion: reduce)').matches",
    "if (!pRM) {",
    "  // all GSAP ScrollTrigger animations here",
    "}",
])

h2("Lenis Final Config")
code_block([
    "const lenis = new Lenis({",
    "  duration: 1.2,  // tune in Phase 2 -- range 0.8-1.4",
    "  easing: (t) => Math.min(1, 1.001 - Math.pow(2, -10 * t)),",
    "  smoothWheel: true,",
    "  wheelMultiplier: 0.8,",
    "  touchMultiplier: 1.5,",
    "  infinite: false,",
    "})",
])

h2("Performance Checklist")
for item in [
    "Preload hero screenshot: <link rel='preload' as='image' href='/audit-result-screenshot.webp'>",
    "Google Fonts: display=swap already set in URL",
    "Geist: load weights 400, 700, 800 only",
    "Instrument Serif: load italic variant only (ital@0,1 -- keep 0 if needed for fallback)",
    "JetBrains Mono: load weight 400 only",
    "All images: loading='lazy' except hero screenshot (loading='eager')",
    "Marquee: will-change: transform on .marquee-track",
]:
    bullet(item)
divider()

# ─────────────────────────────────────────────────────────────────────────────
# FINAL NOTES
# ─────────────────────────────────────────────────────────────────────────────
h1("FINAL NOTES")
body("Build in the order above. Do not jump to a later section before its predecessors are solid. "
     "The design language compounds -- a weak Phase 1 makes every stage harder.")
body("This is a frontend prototype only. The URLBar is decorative -- no POST needed. "
     "All testimonials, stats, and prices are real and verified against the live site at pagegains.com.")
body("When in doubt, push harder on type size and shadow. "
     "This redesign should feel like it is daring someone to scroll past it.")

h2("Key Reference Files")
kv("Original copy",   "original/index.html -- source of truth for all copy and FAQ answers")
kv("Brand guide",     "brand_assets/BRAND.md")
kv("Design tokens",   "brand_assets/tokens.css -- copy to src/tokens.css")
kv("Live site",       "https://pagegains.com/")
kv("Design system",   "brand_assets/design-system.html")

# ─────────────────────────────────────────────────────────────────────────────
out = "/sessions/serene-lucid-goodall/mnt/Pagegains/PageGains_Redesign_Plan.docx"
doc.save(out)
print(f"Saved: {out}")
